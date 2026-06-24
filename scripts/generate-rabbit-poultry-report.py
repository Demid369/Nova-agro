#!/usr/bin/env python3
"""Generate DOCX report: TEO RAG test + rabbit to poultry replacement analysis."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reports" / "TEO_RAG_отчет_кролики_птица.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("Отчёт: тест TEO RAG и замена кролиководства на птицеводство", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Проект «МОЯ МЕЧТА» | {date.today().isoformat()}")
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph()

    # 1
    add_heading(doc, "1. Тестовый запрос и проверка системы", 1)
    add_para(doc, "Тестовый запрос:", bold=True)
    add_para(doc, 'python scripts/teo-query.py "кролиководство связи комбикорм экспорт NPV" --mode hybrid --synthesize')

    add_table(
        doc,
        ["Компонент", "Результат", "Оценка"],
        [
            ["Роутер", "hybrid", "Корректно — и связи, и факты"],
            ["Graphify", "115 узлов, старт: NPV, Кролиководство", "Работает"],
            ["Vector", "KPI-таблица, маркетинг 4.2, кормление", "Работает"],
            ["Synthesis", "Extractive, 6 фрагментов с citations", "Работает"],
            ["Validation", "valid: true, 7 чисел, 5 claims", "Работает"],
        ],
    )

    add_para(doc, "Ключевые факты из ответа (подтверждены corpus):", bold=True)
    add_bullet(doc, "NPV кролиководства: 2 779 519 тыс. руб., IRR 15,19%, payback 84 мес. (00-summary.md)")
    add_bullet(doc, "CAPEX блока: 12 млрд руб.")
    add_bullet(doc, "Производство: 7 000 т мяса/год, оборудование Meneghin Srl")
    add_bullet(doc, "Экспорт: 100% готовой продукции — Азия, СНГ, арабские страны (04-rynok, 4.2)")
    add_para(doc, "Вывод: гибридный RAG отработал корректно — граф дал кластер, вектор подтянул KPI, валидатор не пропустил неподтверждённые числа.")

    # 2
    add_heading(doc, "2. Уточнение: «зайцы» vs «кролики»", 1)
    add_para(doc, "В ТЭО проектный блок — кролиководство, не звероводство (зайцы).")
    add_bullet(doc, "Зайцы встречаются только в рыночной статистике FAOSTAT («мясо кролика и зайца») — классификация ВПС, не проектный актив.")
    add_bullet(doc, "Весь CAPEX, NPV, технология и оборудование завязаны на кроликов.")
    add_para(doc, "Для замены на птицу речь идёт о блоке «Кролиководство» в docs/graphify-corpus/00-summary.md и 01-vvedenie-i-resume.md.")

    # 3
    add_heading(doc, "3. Связи кролиководства в графе", 1)
    add_para(doc, 'graphify explain "Кролиководство" — 20 рёбер, degree 20.')

    add_heading(doc, "3.1. Прямые связи (ядро проекта)", 2)
    add_table(
        doc,
        ["Связь", "Тип", "Смысл"],
        [
            ["ООО «МОЯ МЕЧТА»", "conceptually_related_to", "Блок ядра"],
            ["Кролиководческая ферма", "shares_data_with", "81 цех, фаза I"],
            ["Животноводство", "shares_data_with", "Соседний мясной блок"],
            ["Мясо кролика", "shares_data_with", "Продукт 7 000 т/год"],
            ["Кроличий навоз", "conceptually_related_to", "→ биогаз / удобрения"],
            ["Мех и шкурка кролика", "conceptually_related_to", "Побочный продукт"],
        ],
    )

    add_heading(doc, "3.2. Финансы", 2)
    add_table(
        doc,
        ["Связь", "Значение"],
        [
            ["NPV кролиководство 2,78 млрд", "rationale_for"],
            ["NPV, IRR, Payback, CAPEX", "shares_data_with"],
            ["Ставка дисконтирования 10%", "rationale_for"],
            ["Горизонт расчёта 10 лет", "rationale_for"],
        ],
    )

    add_heading(doc, "3.3. Технологии и поставщики", 2)
    add_bullet(doc, "Meneghin Srl — оборудование фермы (69 цехов интенсивного выращивания)")
    add_bullet(doc, "SINT Technologies — бойня 2 400 голов/час на площадке")
    add_bullet(doc, "F.lli FRAGOLA — комбикорм 262 800 т/год")
    add_bullet(doc, "Gozzini — через path к Комбикормовому заводу")

    add_heading(doc, "3.4. Цепочки path", 2)
    add_para(doc, "Кролиководство → Meneghin Srl → Gozzini → Комбикормовый завод (3 hops)")
    add_para(doc, "Кролиководство → ООО «МОЯ МЕЧТА» → Убойный цех (2 hops)")
    add_para(doc, "Кролиководство → … → Биогазовая установка (3 hops)")

    add_heading(doc, "3.5. Параметры блока (из введения)", 2)
    add_bullet(doc, "До 6 млн голов, 7 000 т мяса/год")
    add_bullet(doc, "Навоз 43 800 т → удобрение")
    add_bullet(doc, "Цена внутренний рынок: 560 руб./кг")
    add_bullet(doc, "Солнечные панели на крышах кролиководческих ферм")

    # 4
    add_heading(doc, "4. Что есть по птице в корпусе сейчас", 1)
    add_para(doc, "Птицеводство как производственный блок проекта — отсутствует.", bold=True)
    add_para(doc, "В 00-summary.md блоки: кролики, теплицы, КРС/МРС, рыба, масложировой. Птицы нет.")
    add_para(doc, "Птица присутствует только как:", bold=True)
    add_bullet(doc, "Рыночный контекст (04-rynok) — конкуренция с кроликом, HS 0207, бройлеры, экспорт в ОАЭ/Китай")
    add_bullet(doc, "Общая статистика РФ — 5,27 млн т мяса птицы в 2024")
    add_bullet(doc, "Комбикорм (03-proizvodstvo) — рецепты/премиксы «для птиц», мясоперьевая мука")
    add_bullet(doc, "AIA — Associazione Italiana Allevatori (крупный рогатый скот), не птицеводство")

    add_para(doc, 'Тест: teo-query "заменить кролиководство на птицеводство" вернул кормление кроликов — корректное поведение, т.к. сценария замены в corpus нет.')

    # 5
    add_heading(doc, "5. Как сделать замену кроликов → птица", 1)

    add_heading(doc, "5.1. Правка корпуса", 2)
    add_table(
        doc,
        ["Файл", "Что менять"],
        [
            ["00-summary.md", "Блок Кролиководство → Птицеводство, KPI, CAPEX, т/год"],
            ["01-vvedenie-i-resume.md", "Фаза I, оборудование, бойня, генетика"],
            ["05-finansy-i-byudzhet.md", "NPV/IRR/payback нового блока"],
            ["04-rynok-i-analitika.md", "4.2 — цены, экспорт HS 0207"],
            ["docs/teo/", "Новые разделы по видам птицы"],
        ],
    )

    add_heading(doc, "5.2. Пересборка артефактов", 2)
    add_para(doc, "python scripts/build-teo-vector-index.py")
    add_para(doc, "uv tool run --from graphifyy python scripts/build-full-teo-graph.py")
    add_para(doc, "uv tool run --from graphifyy python scripts/build-smart-semantic-graph.py")
    add_para(doc, "uv tool run --from graphifyy python scripts/label-teo-communities.py")

    add_heading(doc, "5.3. Обновление онтологии графа", 2)
    add_bullet(doc, "Убрать/заменить: Кролиководство, Meneghin Srl (если не остаётся)")
    add_bullet(doc, "Добавить: Птицеводство, Бройлер, Утиное мясо, Перепел, Индейка")
    add_bullet(doc, "Гиперребро: Птица → бойня → переработка → экспорт HS 0207")

    add_heading(doc, "5.4. Что переносится «как есть»", 2)
    add_table(
        doc,
        ["Связь кроликов", "Аналог для птицы"],
        [
            ["Комбикормовый завод", "Тот же завод, другие рецепты (премикс для птиц уже в 03-proizvodstvo)"],
            ["Убойный цех / SINT", "Бойня птицы (другая линия)"],
            ["Биогаз ← навоз", "Птичий помёт → биогаз (другая инженерия)"],
            ["Кровяная мука", "Мясоперьевая мука"],
            ["Экспорт СНГ/ОАЭ", "HS 0207 — больше данных в 04-rynok"],
            ["Солнечные панели", "Переносится на птичники"],
        ],
    )

    add_heading(doc, "5.5. Что не переносится напрямую", 2)
    add_table(
        doc,
        ["Кролики", "Птица"],
        [
            ["Meneghin Srl (клеточное кролиководство)", "Другое оборудование (инкубаторы, бройлерные цеха)"],
            ["6 млн голов, 81 цех", "Другая плотность и CAPEX"],
            ["Мех/шкурка кролика", "Перья → мясоперьевая мука"],
            ["NPV 2,78 млрд / IRR 15,19%", "Новый финмодель"],
            ["ANCI (итальянские кролиководы)", "Линии Ross/Cobb, утка, перепёл"],
        ],
    )

    # 6
    add_heading(doc, "6. Рекомендация", 1)
    add_para(doc, "Технически наша настройка (Graph + Vector + RAG) поддерживает замену: правка md → reindex → regraph → запросы.")
    add_para(doc, "По смыслу ТЭО замена — не переименование узла, а новый производственный блок с пересчётом CAPEX, NPV, поставщиков и экспортной номенклатуры.")
    add_para(doc, "Прагматичный путь: не удалять кроликов, а добавить птицеводство как отдельный блок — в 04-rynok уже есть рынок HS 0207 и экспорт в ОАЭ (+220% в 2024). Граф покажет два параллельных кластера на общей инфраструктуре (комбикорм, бойня, биогаз, экспорт).", bold=True)

    # 7
    add_heading(doc, "7. Итоговая таблица", 1)
    add_table(
        doc,
        ["Вопрос", "Ответ"],
        [
            ["Система работает?", "Да — тестовый запрос прошёл, validation OK"],
            ["Связи с зайцами?", "В проекте нет; есть кролиководство + FAOSTAT в аналитике"],
            ["Замена на птицу в RAG сейчас?", "Невозможна без правки корпуса"],
            ["Как сделать замену?", "Правка graphify-corpus + teo → reindex + regraph"],
            ["Что сохранить?", "Комбикорм, бойня, биогаз, экспортный кластер"],
            ["Что пересчитать?", "CAPEX, NPV, оборудование, маркетинг 4.2"],
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Сгенерировано: TEO RAG (Nova-agro) | scripts/generate-rabbit-poultry-report.py")
    fr.italic = True
    fr.font.size = Pt(9)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
