#!/usr/bin/env python3
"""Validate RAG corpus vs original TEO for rabbit-farming facts."""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from krolikovodstvo_inventory import (  # noqa: E402
    INVENTORY_DIR,
    RABBIT_RE,
    REPORTS_DIR,
    load_registry,
    read_text,
    resolve_path,
)

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
JSON_OUT = REPORTS_DIR / "rag-validation.json"


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.lower().replace("\u00a0", " "))


def fact_present(text: str, *needles: str) -> bool:
    if not text:
        return False
    t = norm(text)
    t_compact = t.replace(" ", "")
    for n in needles:
        nlow = n.lower()
        if nlow not in t and nlow.replace(" ", "") not in t_compact:
            return False
    return True


def check_canonical_facts(registry: dict) -> list[dict]:
    facts = registry["canonical_facts"]
    checks = [
        ("output_t_per_year", ["7 000", "7000"], "7 000 т/год"),
        ("capex_bln_rub", ["12"], "CAPEX 12 млрд"),
        ("npv_thousand_rub", ["2779519", "2 779 519", "2 779"], "NPV 2 779 519"),
        ("irr_pct", ["15,19", "15.19"], "IRR 15,19%"),
        ("payback_months", ["84"], "payback 84 мес"),
        ("equipment", ["meneghin"], "Meneghin Srl"),
        ("slaughter_heads_per_hour", ["2400"], "SINT 2400 г/ч"),
        ("manure_t_per_year", ["43 800", "43800"], "навоз 43 800 т"),
        ("genetics", ["anci"], "ANCI"),
    ]
    layers = {
        "source_teo": read_text("docs/teo/01-23-потенциал-существующий-и-прогнозируемый.md")
        + read_text("docs/teo/124-кролики-особенности-содержания-и-разведения.md")
        + read_text("docs/graphify-corpus/01-vvedenie-i-resume.md"),
        "corpus": read_text("docs/graphify-corpus/00-summary.md")
        + read_text("docs/graphify-corpus/01-vvedenie-i-resume.md"),
        "kpi_json": read_text("teo-rag-out/kpi.json"),
    }
    bm25 = read_text("teo-rag-out/bm25-index.json")
    rows: list[dict] = []
    for key, needles, label in checks:
        row: dict[str, Any] = {"fact": key, "label": label, "layers": {}}
        for layer, text in layers.items():
            row["layers"][layer] = fact_present(text, *needles)
        row["layers"]["bm25_index"] = fact_present(bm25, *needles)
        # OK если есть в corpus или kpi (RAG рабочий слой)
        row["ok"] = row["layers"]["corpus"] or row["layers"]["kpi_json"]
        rows.append(row)
    return rows


def file_coverage(registry: dict) -> list[dict]:
    rows: list[dict] = []
    all_files = registry.get("all_source_files", {})
    chunk_text = read_text("teo-rag-out/chunk-index.json")
    for group, files in all_files.items():
        for rel in files:
            path = resolve_path(rel)
            src = read_text(rel) if path.exists() else ""
            src_hits = len(RABBIT_RE.findall(src)) if src else 0
            in_chunk = rel.replace("\\", "/") in chunk_text if src else False
            is_system = group == "system_rebuild_after_replace" or group == "structured"
            ok = path.exists() and (is_system or src_hits == 0 or in_chunk)
            rows.append(
                {
                    "group": group,
                    "file": rel,
                    "exists": path.exists(),
                    "rabbit_mentions_source": src_hits,
                    "in_chunk_index": in_chunk,
                    "ok": ok,
                }
            )
    return rows


def rag_gaps() -> list[dict]:
    """Known ingestion gaps for rabbit content."""
    gaps = [
        {
            "id": "teo-125-excluded",
            "severity": "medium",
            "detail": "docs/teo/125-табл-141-рецепты-комбикормов-для-кроликов.md "
            "не попадает в vector index: is_trade_stat_file() фильтрует «табл-N».",
            "mitigation": "Контент дублируется в docs/graphify-corpus/03-proizvodstvo (Табл. 141).",
        },
        {
            "id": "tonnage-graph-mode",
            "severity": "low",
            "detail": "Запрос «сколько тонн крольчатины» без KPI-слова уходит в graph → "
            "мировой рынок, а не 7 000 т проекта.",
            "mitigation": "Использовать KPI fast path: «кролиководство NPV» / «7 000 т кролик».",
        },
        {
            "id": "irr-in-kpi-only",
            "severity": "low",
            "detail": "IRR 15,19% для кроликов — в 00-summary и kpi.json; "
            "в 01-vvedenie может отсутствовать дословно.",
            "mitigation": "KPI-слой teo-query покрывает NPV/IRR запросы.",
        },
    ]
    return gaps


def chunk_block_stats() -> dict:
    text = read_text("teo-rag-out/chunk-index.json")
    if not text:
        return {"rabbit_block_chunks": 0, "total_chunks": 0}
    data = json.loads(text)
    chunks = data if isinstance(data, list) else data.get("chunks", [])
    rabbit = sum(1 for c in chunks if c.get("block") == "кролиководство")
    return {"rabbit_block_chunks": rabbit, "total_chunks": len(chunks)}


def run_rag_queries(registry: dict) -> list[dict]:
    queries = registry.get("rag_validation_queries", [])
    rows: list[dict] = []
    teo_query = ROOT / "scripts" / "teo-query.py"
    if not teo_query.exists():
        return rows
    for item in queries:
        q = item["query"]
        expect = item.get("expect", [])
        try:
            proc = subprocess.run(
                [sys.executable, str(teo_query), q, "--mode", "hybrid"],
                capture_output=True,
                text=True,
                timeout=120,
                cwd=str(ROOT),
            )
            answer = (proc.stdout or "") + (proc.stderr or "")
        except Exception as exc:
            answer = str(exc)
        hit = any(e.lower() in answer.lower() for e in expect) if expect else bool(answer.strip())
        rows.append({"query": q, "expect": expect, "ok": hit, "answer_preview": answer[:500]})
    return rows


def teo_vs_corpus_gaps(registry: dict) -> list[dict]:
    """Primary teo files: rabbit mentions should appear in corpus."""
    gaps: list[dict] = []
    primary = registry.get("all_source_files", {}).get("source_teo_primary", [])
    corpus = (
        read_text("docs/graphify-corpus/01-vvedenie-i-resume.md")
        + read_text("docs/graphify-corpus/03-proizvodstvo-i-tehnologii.md")
        + read_text("docs/graphify-corpus/04-rynok-i-analitika.md")
    )
    for rel in primary:
        src = read_text(rel)
        if not src or not RABBIT_RE.search(src):
            continue
        # dedicated files mapped to corpus sections
        name = Path(rel).name
        in_corpus = name.split("-", 1)[-1][:20] in corpus or RABBIT_RE.search(corpus) is not None
        sample = RABBIT_RE.findall(src)[:3]
        gaps.append(
            {
                "teo_file": rel,
                "teo_mentions": len(RABBIT_RE.findall(src)),
                "reflected_in_corpus": in_corpus,
                "note": "124/125/08 должны быть в 03-proizvodstvo или 04-rynok",
            }
        )
    return gaps


def build_report(run_queries: bool = True) -> dict:
    registry = load_registry()
    facts = check_canonical_facts(registry)
    files = file_coverage(registry)
    report = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "project": registry["meta"]["project"],
        "block": registry["meta"]["block_id"],
        "canonical_facts_check": facts,
        "facts_ok": sum(1 for f in facts if f["ok"]),
        "facts_total": len(facts),
        "file_coverage": files,
        "files_ok": sum(1 for f in files if f["ok"]),
        "files_total": len(files),
        "chunk_stats": chunk_block_stats(),
        "teo_corpus_gaps": teo_vs_corpus_gaps(registry),
        "rag_gaps": rag_gaps(),
    }
    if run_queries:
        report["rag_queries"] = run_rag_queries(registry)
        report["queries_ok"] = sum(1 for q in report["rag_queries"] if q["ok"])
        report["queries_total"] = len(report["rag_queries"])
    return report


def write_docx(report: dict) -> Path:
    out = REPORTS_DIR / "rag-validation.docx"
    doc = Document()
    t = doc.add_heading("Сверка RAG vs исходный ТЭО (кролиководство)", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(date.today().isoformat())
    r.italic = True
    r.font.size = Pt(10)

    doc.add_heading("Итог", level=1)
    doc.add_paragraph(
        f"Факты KPI: {report['facts_ok']}/{report['facts_total']} | "
        f"Файлы в индексе: {report['files_ok']}/{report['files_total']} | "
        f"Чанков block=кролиководство: {report['chunk_stats'].get('rabbit_block_chunks', 0)}"
    )
    if "queries_ok" in report:
        doc.add_paragraph(f"RAG-запросы: {report['queries_ok']}/{report['queries_total']}")

    doc.add_heading("Канонические факты по слоям", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = ["Факт", "corpus", "source_teo", "kpi", "bm25"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    for row in report["canonical_facts_check"]:
        cells = table.add_row().cells
        cells[0].text = row["label"]
        layers = row["layers"]
        for i, key in enumerate(["corpus", "source_teo", "kpi_json", "bm25_index"], start=1):
            cells[i].text = "OK" if layers.get(key) else "—"

    doc.add_heading("Покрытие файлов", level=1)
    for f in report["file_coverage"]:
        if f["group"] == "system_rebuild_after_replace":
            continue
        status = "OK" if f["ok"] else "ПРОВЕРИТЬ"
        doc.add_paragraph(
            f"[{status}] {f['file']} — упоминаний: {f['rabbit_mentions_source']}, "
            f"в chunk-index: {'да' if f['in_chunk_index'] else 'нет'}"
        )

    doc.add_heading("Известные пробелы RAG", level=1)
    for g in report.get("rag_gaps", []):
        doc.add_paragraph(f"[{g['severity']}] {g['detail']}")
        doc.add_paragraph(f"→ {g['mitigation']}", style="List Bullet")

    if report.get("rag_queries"):
        doc.add_heading("Тестовые запросы teo-query", level=1)
        for q in report["rag_queries"]:
            doc.add_paragraph(f"{'OK' if q['ok'] else 'FAIL'}: {q['query']}")

    doc.save(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", action="store_true", help="Also write rag-validation.docx")
    parser.add_argument("--no-queries", action="store_true", help="Skip live teo-query subprocess")
    args = parser.parse_args()

    report = build_report(run_queries=not args.no_queries)
    JSON_OUT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {JSON_OUT}")
    print(
        f"Facts: {report['facts_ok']}/{report['facts_total']} | "
        f"Files: {report['files_ok']}/{report['files_total']}"
    )
    if args.docx:
        p = write_docx(report)
        print(f"Wrote {p}")


if __name__ == "__main__":
    main()
