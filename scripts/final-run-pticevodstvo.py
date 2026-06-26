#!/usr/bin/env python3
"""Final run: generate DOCX + validate registry + KPI smoke check on master draft."""

from __future__ import annotations

import re
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
MASTER = ROOT / "docs/inventory/pticevodstvo/docx/00-master-teo-pticevodstvo-draft.docx"
REPORT = ROOT / "docs/inventory/pticevodstvo/reports/final-run-jun2026.md"

sys.path.insert(0, str(ROOT / "scripts"))
from pticevodstvo_inventory import validate_registry, load_registry  # noqa: E402


def doc_full_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def kpi_ok(text: str) -> list[tuple[str, bool]]:
    flat = text.replace("\u00a0", " ")
    return [
        ("CAPEX 12 000", bool(re.search(r"12\s*000", flat))),
        ("Revenue 5 559", bool(re.search(r"5\s*559", flat))),
        ("EBITDA 2 216", bool(re.search(r"2\s*216", flat))),
        ("NPV +2 253", bool(re.search(r"2\s*253", flat))),
        ("IRR 12,8%", bool(re.search(r"12[,.]8", flat))),
        ("476 FTE", bool(re.search(r"\b476\b", flat))),
        ("118 птичников", bool(re.search(r"\b118\b", flat))),
        ("T004 ВСЕГО 72 158 537", "72158537" in flat.replace(" ", "")),
        ("T241 Птицеводство", "Птицеводство" in flat),
        ("DCF 16 лет", bool(re.search(r"16\s*лет", flat))),
    ]


def main() -> int:
    gen = subprocess.run(
        [sys.executable, str(ROOT / "scripts/generate-pticevodstvo-docx.py")],
        cwd=ROOT,
        capture_output=True,
        text=True,
    )
    if gen.returncode != 0:
        print(gen.stdout, gen.stderr, sep="\n", file=sys.stderr)
        return gen.returncode

    errors = validate_registry(load_registry())
    if errors:
        print("registry validation FAILED:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        return 1

    if not MASTER.exists():
        print(f"missing {MASTER}", file=sys.stderr)
        return 1

    text = doc_full_text(MASTER)
    words = len(text.split())
    checks = kpi_ok(text)
    failed = [name for name, ok in checks if not ok]

    print(f"final run OK @ {datetime.now(timezone.utc).isoformat()}")
    print(f"  master: {words} words, {MASTER}")
    for name, ok in checks:
        print(f"  {'✓' if ok else '✗'} {name}")

    if failed:
        print(f"KPI check FAILED: {', '.join(failed)}", file=sys.stderr)
        return 1

    print(f"  report: {REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
