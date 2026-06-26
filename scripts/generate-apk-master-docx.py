#!/usr/bin/env python3
"""Generate full APK master DOCX: standalone poultry + corpus (patched) + appendix A trade tables."""

from __future__ import annotations

import argparse
import importlib.util
import json
import re
import sys
from pathlib import Path

import yaml
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

_gen_spec = importlib.util.spec_from_file_location(
    "generate_pticevodstvo_docx",
    SCRIPTS / "generate-pticevodstvo-docx.py",
)
_gen = importlib.util.module_from_spec(_gen_spec)
assert _gen_spec and _gen_spec.loader
_gen_spec.loader.exec_module(_gen)

build_master_docx = _gen.build_master_docx
write_markdown_body = _gen.write_markdown_body
add_meta = _gen.add_meta

from pticevodstvo_inventory import (  # noqa: E402
    DOCX_DIR,
    assert_registry_valid,
    load_master_assembly,
    load_registry,
    read_text,
    write_docx_audit,
)

APK_ASSEMBLY_PATH = ROOT / "docs/inventory/pticevodstvo/pipeline/assembly-apk-full.yaml"
PATCH_RULES_PATH = ROOT / "docs/inventory/pticevodstvo/pipeline/corpus-patch-rules.yaml"
MANIFEST_PATH = ROOT / "docs/teo-tables/manifest.json"
CORPUS_MAX_CHARS = 500_000
APPENDIX_TABLE_LIMIT = 0  # 0 = all reference_trade


def load_apk_assembly() -> dict:
    return yaml.safe_load(APK_ASSEMBLY_PATH.read_text(encoding="utf-8"))


def load_patch_rules() -> dict:
    if not PATCH_RULES_PATH.exists():
        return {}
    return yaml.safe_load(PATCH_RULES_PATH.read_text(encoding="utf-8")) or {}


def filter_corpus_text(text: str, exclude_headings: list[str], line_patterns: list[str]) -> str:
    if not text:
        return text

    lines = text.splitlines()
    out: list[str] = []
    skip = False
    heading_re = re.compile(r"^#\s+")

    for line in lines:
        stripped = line.strip()
        if heading_re.match(stripped):
            title = stripped.lstrip("#").strip()
            skip = any(ex.lower() in title.lower() for ex in exclude_headings)
        if skip:
            continue
        if any(re.search(p, line, re.IGNORECASE) for p in line_patterns):
            continue
        out.append(line)

    return "\n".join(out)


def add_corpus_footnote(doc: Document) -> None:
    p = doc.add_paragraph(
        "Corpus baseline (APK blocks 2–5 as-is). Блок 1 I-фазы — птицеводство «Нова-Агро» "
        "(review-pass R-01…R-05). Секции кролиководства вырезаны patch-rules."
    )
    for run in p.runs:
        run.italic = True


def append_section(doc: Document, title: str, section_ref: str | None, note: str | None) -> None:
    doc.add_page_break()
    heading = title if not section_ref else f"{title} ({section_ref})"
    doc.add_heading(heading, level=1)
    if note:
        p = doc.add_paragraph(note)
        for run in p.runs:
            run.italic = True


def build_apk_master_docx(registry: dict, apk_asm: dict, patch_rules: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    meta = apk_asm.get("meta", {})
    out_name = meta.get("output_docx", "00-apk-master-teo-full-draft.docx")
    out = DOCX_DIR / out_name

    standalone = load_master_assembly()
    standalone_doc_path = build_master_docx(registry, standalone)
    doc = Document(str(standalone_doc_path))

    line_patterns = patch_rules.get("line_skip_patterns") or []

    doc.add_page_break()
    doc.add_heading("Часть II — Corpus baseline (blocks 2–5)", level=0)
    add_meta(doc, "Patch-rules: corpus-patch-rules.yaml", block_label="APK corpus")

    for sec in apk_asm.get("corpus_sections", []):
        append_section(doc, sec.get("title", sec.get("id", "?")), sec.get("section_ref"), sec.get("note"))
        add_corpus_footnote(doc)

        patch = sec.get("patch") or {}
        exclude = patch.get("exclude_headings") or []

        for src in sec.get("sources", []):
            rel = src.get("file", "")
            if not rel:
                continue
            body = read_text(rel)
            body = filter_corpus_text(body, exclude, line_patterns)
            if len(body) > CORPUS_MAX_CHARS:
                body = body[:CORPUS_MAX_CHARS] + "\n\n[… corpus truncated …]"
            doc.add_heading(rel, level=2)
            write_markdown_body(doc, body, rel)

    for sec in apk_asm.get("appendix_sections", []):
        append_section(doc, sec.get("title", "?"), sec.get("section_ref"), sec.get("note"))
        manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
        tables = [t for t in manifest.get("tables", []) if t.get("category") == sec.get("category")]
        if APPENDIX_TABLE_LIMIT:
            tables = tables[:APPENDIX_TABLE_LIMIT]

        doc.add_paragraph(f"Таблиц reference_trade: {len(tables)}")
        for t in tables:
            path = t.get("path_all") or t.get("path_critical")
            if not path:
                continue
            body = read_text(path)
            if not body:
                continue
            doc.add_heading(f"{t.get('id', '?')} — {t.get('title', '')[:80]}", level=3)
            write_markdown_body(doc, body, path)

    doc.save(out)
    return out


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--appendix-limit", type=int, default=0, help="Limit trade tables (0=all)")
    args = parser.parse_args()

    global APPENDIX_TABLE_LIMIT
    APPENDIX_TABLE_LIMIT = args.appendix_limit

    registry = load_registry()
    try:
        assert_registry_valid(registry)
    except Exception as exc:
        print(f"registry validation FAILED:\n{exc}", file=sys.stderr)
        return 1

    apk_asm = load_apk_assembly()
    patch_rules = load_patch_rules()
    audit_path = write_docx_audit(registry)
    print(f"Wrote {audit_path}")

    out = build_apk_master_docx(registry, apk_asm, patch_rules)
    print(f"APK master: {out}")
    print(f"  standalone sections: {len(load_master_assembly().get('sections', []))}")
    print(f"  corpus sections: {len(apk_asm.get('corpus_sections', []))}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
