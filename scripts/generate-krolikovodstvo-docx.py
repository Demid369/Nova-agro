#!/usr/bin/env python3
"""Generate DOCX files per rabbit-farming theme from docs/inventory/krolikovodstvo/registry.yaml."""

from __future__ import annotations

import argparse
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from krolikovodstvo_inventory import (
    DOCX_DIR,
    ThemeChunk,
    assert_registry_valid,
    canonical_fact_refs,
    gather_theme_content,
    load_registry,
    write_docx_audit,
)

MAX_BODY_CHARS = 120_000


def add_meta(doc: Document, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Проект «МОЯ МЕЧТА» | baseline кролиководство | {date.today().isoformat()}")
    r.italic = True
    r.font.size = Pt(10)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
    doc.add_paragraph()


def add_ref_paragraph(doc: Document, ref: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"↳ {ref}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)


def write_theme_chunks(doc: Document, chunks: list[ThemeChunk]) -> None:
    for chunk in chunks:
        doc.add_heading(chunk.header, level=2)
        if chunk.refs:
            doc.add_paragraph("Ссылки на исходный ТЭО (файл:строка):")
            for ref in chunk.refs:
                add_ref_paragraph(doc, ref)

        for frag in chunk.fragments:
            if len(chunk.fragments) > 1:
                doc.add_heading(f"Фрагмент {frag.ref}", level=3)
            text = frag.text
            if len(text) > MAX_BODY_CHARS:
                text = text[:MAX_BODY_CHARS] + "\n\n[… обрезано …]"
            for para in text.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                if para.startswith("#"):
                    level = len(para) - len(para.lstrip("#"))
                    title = para.lstrip("#").strip()
                    doc.add_heading(title, level=min(level + 1, 4))
                else:
                    if frag.line_start == frag.line_end and not para.startswith("↳"):
                        p = doc.add_paragraph()
                        anchor = p.add_run(f"[L{frag.line_start}] ")
                        anchor.font.size = Pt(8)
                        anchor.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        p.add_run(para)
                    else:
                        doc.add_paragraph(para)


def build_theme_docx(theme_id: str, theme: dict, registry: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCX_DIR / theme["docx"]
    doc = Document()
    title = doc.add_heading(f"{theme_id}. {theme['name']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, theme.get("description", ""))

    doc.add_heading("Действие при замене блока", level=1)
    doc.add_paragraph(f"action_on_replace: {theme.get('action_on_replace', 'replace')}")

    if theme_id == "T13":
        _write_t13_body(doc, registry)
    else:
        chunks = gather_theme_content(theme)
        if not chunks:
            doc.add_paragraph("Нет извлечённых фрагментов — проверьте registry.yaml.")
        else:
            doc.add_heading("Содержание по источникам", level=1)
            write_theme_chunks(doc, chunks)

    doc.save(out)
    return out


def _write_t13_body(doc: Document, registry: dict) -> None:
    doc.add_heading("Правило", level=1)
    doc.add_paragraph(
        "В DOCX-темах canonical = graphify-corpus. Файлы teo/124 и teo/02 — дубли, "
        "намеренно не включены в тематические DOCX (остаются в RAG и all_source_files)."
    )

    rows = registry.get("duplicates_map", [])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["teo файл", "corpus файл", "статус", "в DOCX", "примечание"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row.get("teo", "")
        cells[1].text = (row.get("corpus") or "—") + (
            f" ({row.get('corpus_anchor', '')})" if row.get("corpus_anchor") else ""
        )
        cells[2].text = row.get("status", "")
        cells[3].text = "да" if row.get("in_docx") else "нет"
        cells[4].text = row.get("note", "")

    doc.add_paragraph()
    doc.add_heading("Легенда статусов", level=1)
    for line in [
        "duplicate — полный или почти полный дубль corpus",
        "partial — пересечение по части разделов",
        "unique — только в teo, нет аналога в corpus",
        "episodic — эпизодическое упоминание, не проектный блок",
    ]:
        doc.add_paragraph(line, style="List Bullet")


def build_index_docx(registry: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCX_DIR / "00-index-krolikovodstvo.docx"
    doc = Document()
    t = doc.add_heading("Инвентарь кролиководства — оглавление", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, "Все направления для будущей замены в новом ТЭО")

    doc.add_heading("Канонические KPI — ссылки на строки", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[0].cells[2].text = "Источник (файл:строка)"
    for row in canonical_fact_refs(registry):
        cells = table.add_row().cells
        cells[0].text = row["fact"]
        cells[1].text = row["label"]
        cells[2].text = row["ref"]
    doc.add_paragraph()

    doc.add_heading("DOCX по направлениям", level=1)
    for tid in sorted(k for k in registry.get("themes", {}) if k.startswith("T")):
        th = registry["themes"][tid]
        doc.add_paragraph(f"{tid} — {th['name']}: {th['docx']}", style="List Bullet")
    doc.add_paragraph(
        "00-replace-checklist.docx — только replace/adapt + rebuild (чеклист замены)",
        style="List Bullet",
    )

    doc.add_heading("Все файлы-источники (замена)", level=1)
    all_files = registry.get("all_source_files", {})
    for group, files in all_files.items():
        doc.add_heading(group, level=2)
        for f in files:
            doc.add_paragraph(f, style="List Bullet")

    doc.save(out)
    return out


def build_replace_checklist_docx(registry: dict) -> Path:
    """DOCX with replace/adapt themes only + system rebuild list."""
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCX_DIR / "00-replace-checklist.docx"
    doc = Document()
    title = doc.add_heading("Чеклист замены блока кролиководства", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, "Только темы replace/adapt и артефакты пересборки")

    doc.add_heading("Темы для замены или адаптации", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["ID", "Название", "DOCX", "action_on_replace"]):
        table.rows[0].cells[i].text = h
    for tid in sorted(k for k in registry.get("themes", {}) if k.startswith("T") and k != "T13"):
        th = registry["themes"][tid]
        action = th.get("action_on_replace", "replace")
        if action not in ("replace", "adapt"):
            continue
        cells = table.add_row().cells
        cells[0].text = tid
        cells[1].text = th.get("name", "")
        cells[2].text = th.get("docx", "")
        cells[3].text = action

    doc.add_paragraph()
    doc.add_heading("Пересборка после замены (system_rebuild_after_replace)", level=1)
    rebuild = registry.get("all_source_files", {}).get("system_rebuild_after_replace", [])
    for path in rebuild:
        doc.add_paragraph(path, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("Канонические KPI (контроль)", level=1)
    facts = registry.get("canonical_facts", {})
    for key, val in facts.items():
        doc.add_paragraph(f"{key}: {val}", style="List Bullet")

    doc.save(out)
    return out


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--skip-docx",
        action="store_true",
        help="Validate registry + write docx-audit.json only (no DOCX generation)",
    )
    args = parser.parse_args()

    registry = load_registry()
    try:
        assert_registry_valid(registry)
    except Exception as exc:
        print(f"registry validation FAILED:\n{exc}", file=sys.stderr)
        return 1
    print("registry.yaml OK")

    audit_path = write_docx_audit(registry)
    print(f"Wrote {audit_path}")

    if args.skip_docx:
        return 0

    paths: list[Path] = [
        build_index_docx(registry),
        build_replace_checklist_docx(registry),
    ]
    for tid in sorted(k for k in registry.get("themes", {}) if k.startswith("T")):
        paths.append(build_theme_docx(tid, registry["themes"][tid], registry))
    print(f"Generated {len(paths)} DOCX in docs/inventory/krolikovodstvo/docx")
    for p in paths:
        print(f"  {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
