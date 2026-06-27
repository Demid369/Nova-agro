#!/usr/bin/env python3
"""QA verify baseline→poultry DOCX: drawings, KPI grep, residual «кролик» count."""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

import yaml

ROOT = Path(__file__).resolve().parents[1]
RULES = ROOT / "docs/inventory/pticevodstvo/pipeline/phase1-tables.yaml"
_meta = yaml.safe_load(RULES.read_text(encoding="utf-8"))["meta"]
DEFAULT_BASELINE = ROOT / _meta["baseline"]
DEFAULT_OUTPUT = ROOT / _meta["output"]
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def count_drawings(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        doc = ET.fromstring(zf.read("word/document.xml"))
        return len(doc.findall(".//w:drawing", NS))


def full_text(path: Path) -> str:
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def count_krolik(text: str) -> int:
    return len(re.findall(r"кролик|кроль|кролич", text, re.I))


def kpi_checks(text: str) -> list[tuple[str, bool]]:
    flat = text.replace("\u00a0", " ")
    return [
        ("CAPEX 12 000", bool(re.search(r"12\s*000", flat))),
        ("Revenue 5 559", bool(re.search(r"5\s*559", flat))),
        ("NPV +2 253", bool(re.search(r"2\s*253", flat))),
        ("IRR 12,8%", bool(re.search(r"12[,.]8", flat))),
        ("476 FTE", bool(re.search(r"\b476\b", flat))),
        ("118 птичников", bool(re.search(r"\b118\b", flat))),
        ("SINT 6 000", bool(re.search(r"SINT\s*6\s*000|6000\s*гол", flat, re.I))),
        ("Птицеводство", "Птицеводство" in flat or "птицеводств" in flat.lower()),
    ]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--baseline", default=str(DEFAULT_BASELINE))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    parser.add_argument("--max-krolik", type=int, default=80)
    args = parser.parse_args()

    baseline = Path(args.baseline)
    output = Path(args.output)
    if not output.exists():
        print(f"Missing output: {output}", file=sys.stderr)
        return 1

    b_draw = count_drawings(baseline) if baseline.exists() else -1
    o_draw = count_drawings(output)
    draw_ok = b_draw == o_draw if b_draw >= 0 else True

    text = full_text(output)
    krolik = count_krolik(text)
    checks = kpi_checks(text)
    failed_kpi = [n for n, ok in checks if not ok]

    print(f"Drawings: baseline={b_draw} output={o_draw} {'OK' if draw_ok else 'MISMATCH'}")
    print(f"Residual «кролик/кроль/кролич»: {krolik} (target ≤{args.max_krolik})")
    for name, ok in checks:
        print(f"  {'OK' if ok else 'FAIL'} {name}")

    ok = draw_ok and krolik <= args.max_krolik and not failed_kpi
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
