#!/usr/bin/env python3
"""Build poultry TEO from baseline DOCX: preserve all images, replace block-1 data."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
DEFAULT_RULES = ROOT / "docs/inventory/pticevodstvo/poultry-baseline-replace.yaml"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def load_rules(path: Path) -> dict[str, Any]:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def parse_md_table(md_path: Path) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in md_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return paragraph._p.find(f".//{{{W_NS}}}drawing") is not None


def table_has_drawing(table) -> bool:
    return table._tbl.find(f".//{{{W_NS}}}drawing") is not None


def cell_has_drawing(cell) -> bool:
    return cell._tc.find(f".//{{{W_NS}}}drawing") is not None


def fill_table_from_md(table, rows: list[list[str]], clear_extra: bool = True) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    # expand columns in existing rows if needed (python-docx limited — use min)
    while len(table.rows) < len(rows):
        table.add_row()
    for i, row_data in enumerate(rows):
        for j in range(width):
            val = row_data[j] if j < len(row_data) else ""
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = val
    if clear_extra:
        for i in range(len(rows), len(table.rows)):
            for cell in table.rows[i].cells:
                if cell_has_drawing(cell):
                    continue
                if not cell.text.strip():
                    continue
                cell.text = ""


def fill_table_from_md_skip_image_cells(table, rows: list[list[str]]) -> None:
    """Fill table but never overwrite cells that contain embedded images."""
    if not rows:
        return
    while len(table.rows) < len(rows):
        table.add_row()
    width = max(len(r) for r in rows)
    filled = 0
    for i, row_data in enumerate(rows):
        for j in range(width):
            if j >= len(table.rows[i].cells):
                continue
            cell = table.rows[i].cells[j]
            if cell_has_drawing(cell):
                continue
            val = row_data[j] if j < len(row_data) else ""
            if val:
                cell.text = val
                filled += 1
    return filled


def clear_table_text(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.text = ""


def apply_text_replacements(text: str, rules: list[dict[str, str]]) -> str:
    out = text
    for rule in rules:
        out = re.sub(rule["pattern"], rule["replace"], out, flags=re.IGNORECASE)
    return out


def patch_table14_block1(table, capex_mln: str = "12 000") -> None:
    for row in table.rows:
        for cell in row.cells:
            if re.search(r"кролик", cell.text, re.I):
                cell.text = re.sub(
                    r"кроликовод[^\n]*",
                    f"Птицеводческий комплекс «Нова-Агро» ({capex_mln} млн ₽)",
                    cell.text,
                    flags=re.I,
                )


def map_images(baseline: Path, out_json: Path) -> dict[str, Any]:
    out_json.parent.mkdir(parents=True, exist_ok=True)
    entries: list[dict[str, Any]] = []
    with zipfile.ZipFile(baseline) as zf:
        media_files = sorted(n for n in zf.namelist() if n.startswith("word/media/"))
        doc = ET.fromstring(zf.read("word/document.xml"))
        body = doc.find("w:body", NS)
        tbl_no = 0
        for i, el in enumerate(list(body)):
            tag = el.tag.split("}")[-1]
            if tag == "tbl":
                tbl_no += 1
                continue
            if tag != "p":
                continue
            if not el.findall(".//w:drawing", NS):
                continue
            texts = [t.text or "" for t in el.findall(".//w:t", NS)]
            ctx = "".join(texts).strip()[:200]
            blips = []
            for blip in el.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip", NS):
                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if embed:
                    blips.append(embed)
            entries.append(
                {
                    "body_index": i,
                    "table_before": tbl_no,
                    "context": ctx,
                    "relationship_ids": blips,
                }
            )
        manifest = {
            "source": str(baseline),
            "mapped_at": datetime.now(timezone.utc).isoformat(),
            "media_file_count": len(media_files),
            "inline_image_paragraphs": len(entries),
            "media_files": [Path(m).name for m in media_files],
            "placements": entries,
        }
    out_json.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def extract_media_files(baseline: Path, out_dir: Path) -> int:
    out_dir.mkdir(parents=True, exist_ok=True)
    count = 0
    with zipfile.ZipFile(baseline) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/"):
                continue
            dest = out_dir / Path(name).name
            dest.write_bytes(zf.read(name))
            count += 1
    return count


def insert_footnote_after(doc: Document, para_index: int, text: str) -> None:
    if para_index < len(doc.paragraphs):
        doc.paragraphs[para_index].insert_paragraph_before(text)
    else:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.italic = True


def build_poultry_teo(rules_path: Path, extract_media: bool = False) -> Path:
    rules = load_rules(rules_path)
    meta = rules.get("meta", {})
    baseline = ROOT / meta["baseline"]
    output = ROOT / meta["output"]
    output.parent.mkdir(parents=True, exist_ok=True)

    if not baseline.exists():
        raise FileNotFoundError(f"Baseline not found: {baseline}")

    shutil.copy2(baseline, output)

    media_map_path = ROOT / meta.get("media_map", "docs/inventory/pticevodstvo/media/image-map.json")
    manifest = map_images(baseline, media_map_path)
    print(f"Image map: {len(manifest['placements'])} inline paragraphs, {manifest['media_file_count']} media files")
    print(f"  → {media_map_path}")

    if extract_media:
        media_dir = media_map_path.parent
        n = extract_media_files(baseline, media_dir)
        print(f"Extracted {n} media files → {media_dir}")

    doc = Document(str(output))

    # --- tables ---
    for tab_num_str, md_rel in (rules.get("table_replacements") or {}).items():
        tab_num = int(tab_num_str)
        idx = tab_num - 1
        if idx >= len(doc.tables):
            print(f"WARN: table #{tab_num} missing (doc has {len(doc.tables)})", file=sys.stderr)
            continue
        table = doc.tables[idx]
        md_path = ROOT / md_rel
        if tab_num == 14:
            patch_table14_block1(table)
            print(f"PATCH table #14 (I-фаза row)")
            continue
        if not md_path.exists():
            print(f"WARN: md not found {md_path}", file=sys.stderr)
            continue
        if table_has_drawing(table):
            rows = parse_md_table(md_path)
            n = fill_table_from_md_skip_image_cells(table, rows)
            print(f"REPLACE table #{tab_num} (skip image cells) ← {md_rel} ({n} cells)")
            continue
        rows = parse_md_table(md_path)
        fill_table_from_md(table, rows)
        print(f"REPLACE table #{tab_num} ← {md_rel} ({len(rows)} rows)")

    for tab_num in rules.get("clear_tables") or []:
        idx = tab_num - 1
        if idx >= len(doc.tables):
            continue
        table = doc.tables[idx]
        if table_has_drawing(table):
            print(f"SKIP clear table #{tab_num} (has images)")
            continue
        clear_table_text(table)
        print(f"CLEAR table #{tab_num}")

    # --- paragraphs (no drawing) ---
    preserve_img = rules.get("preserve_paragraphs_with_images", True)
    text_rules = rules.get("text_replacements") or []
    changed = 0
    skipped_img = 0
    for para in doc.paragraphs:
        if preserve_img and paragraph_has_drawing(para):
            skipped_img += 1
            continue
        old = para.text
        if not old.strip():
            continue
        new = apply_text_replacements(old, text_rules)
        if new != old:
            para.text = new
            changed += 1

    # --- table cell text (non-image cells in all tables) ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell_has_drawing(cell):
                    continue
                old = cell.text
                if not old.strip():
                    continue
                new = apply_text_replacements(old, text_rules)
                if new != old:
                    cell.text = new
                    changed += 1

    footnote = rules.get("review_footnote")
    if footnote:
        insert_footnote_after(doc, int(rules.get("insert_review_footnote_after_para", 8)), footnote.strip())

    doc.save(str(output))
    print(f"Text replacements (para+cells): {changed} (skipped {skipped_img} paras with images)")
    print(f"Output: {output} ({output.stat().st_size // 1024 // 1024} MB)")
    return output


def verify_images_preserved(baseline: Path, output: Path) -> bool:
    def count_drawings(path: Path) -> int:
        with zipfile.ZipFile(path) as zf:
            doc = ET.fromstring(zf.read("word/document.xml"))
            return len(doc.findall(".//w:drawing", NS))

    b = count_drawings(baseline)
    o = count_drawings(output)
    ok = b == o
    print(f"Drawings: baseline={b} output={o} {'OK' if ok else 'MISMATCH'}")
    return ok


def main() -> int:
    parser = argparse.ArgumentParser(description="Build poultry TEO from baseline DOCX")
    parser.add_argument("--rules", default=str(DEFAULT_RULES))
    parser.add_argument("--extract-media", action="store_true", help="Also extract word/media/* files")
    parser.add_argument("--skip-phase2", action="store_true", help="Skip §7 narrative + Tab P-141")
    parser.add_argument("--verify-only", action="store_true")
    args = parser.parse_args()

    rules_path = Path(args.rules)
    rules = load_rules(rules_path)
    baseline = ROOT / rules["meta"]["baseline"]
    output = ROOT / rules["meta"]["output"]

    if args.verify_only:
        if not output.exists():
            print(f"Missing output: {output}", file=sys.stderr)
            return 1
        return 0 if verify_images_preserved(baseline, output) else 1

    try:
        out = build_poultry_teo(rules_path, extract_media=args.extract_media)
    except Exception as exc:
        print(f"FAILED: {exc}", file=sys.stderr)
        return 1

    if not verify_images_preserved(baseline, out):
        return 1

    if not args.skip_phase2:
        import importlib.util

        p2_path = SCRIPTS / "apply-teo-poultry-phase2.py"
        spec = importlib.util.spec_from_file_location("apply_teo_poultry_phase2", p2_path)
        p2 = importlib.util.module_from_spec(spec)
        assert spec and spec.loader
        spec.loader.exec_module(p2)
        phase2_yaml = ROOT / "docs/inventory/pticevodstvo/poultry-phase2.yaml"
        p2.apply_phase2(out, phase2_yaml)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
