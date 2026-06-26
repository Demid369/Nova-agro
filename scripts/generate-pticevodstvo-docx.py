#!/usr/bin/env python3
"""Generate DOCX for poultry TEO: themes T01–T12 + master draft from 00-master-assembly.yaml."""

from __future__ import annotations

import argparse
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from pticevodstvo_inventory import (
    DOCX_DIR,
    assert_registry_valid,
    canonical_fact_refs,
    gather_section_sources,
    gather_theme_files,
    load_master_assembly,
    load_registry,
    read_text,
    resolve_path,
    write_docx_audit,
)

MAX_BODY_CHARS = 120_000
TABLE_ROW_RE = re.compile(r"^\|.+\|$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def add_meta(doc: Document, subtitle: str, block_label: str = "птицеводство draft") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Проект «МОЯ МЕЧТА» | {block_label} | {date.today().isoformat()}")
    r.italic = True
    r.font.size = Pt(10)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(10)
    doc.add_paragraph()


def add_ref_paragraph(doc: Document, ref: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"↳ {ref}")
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)


def parse_md_table(lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    rows: list[list[str]] = []
    for line in lines:
        if not TABLE_ROW_RE.match(line.strip()):
            return None
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return None
    # skip separator row |---|
    if all(re.match(r"^:?-+:?$", c.replace(" ", "")) or c == "" for c in rows[1]):
        data_rows = rows[2:]
        headers = rows[0]
    else:
        headers = rows[0]
        data_rows = rows[1:]
    return headers, data_rows


def add_md_table(doc: Document, headers: list[str], data_rows: list[list[str]]) -> None:
    if not headers:
        return
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in data_rows:
        cells = table.add_row().cells
        for i in range(ncols):
            cells[i].text = row[i] if i < len(row) else ""


def write_markdown_body(doc: Document, text: str, source_ref: str | None = None) -> None:
    if source_ref:
        add_ref_paragraph(doc, source_ref)
    if len(text) > MAX_BODY_CHARS:
        text = text[:MAX_BODY_CHARS] + "\n\n[… обрезано …]"

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph("\n".join(code_lines))
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
            i += 1
            continue

        m = HEADING_RE.match(stripped)
        if m:
            level = min(len(m.group(1)) + 1, 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if TABLE_ROW_RE.match(stripped):
            block: list[str] = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i].strip()):
                block.append(lines[i].strip())
                i += 1
            parsed = parse_md_table(block)
            if parsed:
                add_md_table(doc, parsed[0], parsed[1])
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or HEADING_RE.match(nxt) or TABLE_ROW_RE.match(nxt) or nxt.startswith("```"):
                break
            if nxt.startswith("- ") or nxt.startswith("* ") or nxt.startswith(">"):
                break
            para_lines.append(nxt)
            i += 1
        doc.add_paragraph("\n".join(para_lines))


def build_theme_docx(theme_id: str, theme: dict, registry: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    slug = theme_id.lower()
    name_slug = re.sub(r"[^\w\-]+", "-", theme.get("name", theme_id).lower())[:40].strip("-")
    out = DOCX_DIR / f"{slug}-{name_slug}.docx"

    doc = Document()
    title = doc.add_heading(f"{theme_id}. {theme['name']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, theme.get("description", ""))

    doc.add_heading("Действие", level=1)
    doc.add_paragraph(f"action_on_replace: {theme.get('action_on_replace', 'new')}")

    doc.add_heading("Содержание", level=1)
    for rel in gather_theme_files(theme):
        body = read_text(rel)
        if not body:
            doc.add_heading(f"Источник: {rel}", level=2)
            doc.add_paragraph(f"[ФАЙЛ НЕ НАЙДЕН: {rel}]")
            continue
        doc.add_heading(f"Источник: {rel}", level=2)
        write_markdown_body(doc, body, rel)

    doc.save(out)
    return out


def build_index_docx(registry: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCX_DIR / "00-index-pticevodstvo.docx"
    doc = Document()
    t = doc.add_heading("Инвентарь птицеводства — оглавление", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, "Draft — замена блока кролиководства")

    doc.add_heading("Канонические KPI", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Параметр"
    table.rows[0].cells[1].text = "Значение"
    table.rows[0].cells[2].text = "Источник"
    for row in canonical_fact_refs(registry):
        cells = table.add_row().cells
        cells[0].text = row["fact"]
        cells[1].text = row["label"]
        cells[2].text = row["ref"]

    doc.add_heading("Темы T01–T12", level=1)
    for tid in sorted(k for k in registry.get("themes", {}) if k.startswith("T")):
        th = registry["themes"][tid]
        files = ", ".join(gather_theme_files(th))
        doc.add_paragraph(f"{tid} — {th['name']}: {files}", style="List Bullet")

    facts = registry.get("canonical_facts", {})
    doc.add_heading("Critical-таблицы", level=1)
    for key, path in (facts.get("critical_tables") or {}).items():
        doc.add_paragraph(f"{key}: {path}", style="List Bullet")

    doc.add_paragraph("Master draft: 00-master-teo-pticevodstvo-draft.docx", style="List Bullet")
    doc.save(out)
    return out


def build_master_docx(registry: dict, assembly: dict) -> Path:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    meta = assembly.get("meta", {})
    out_name = meta.get("output_docx", "00-master-teo-pticevodstvo-draft.docx")
    out = DOCX_DIR / out_name

    doc = Document()
    title = doc.add_heading(meta.get("title", "Master TEO poultry draft"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_meta(doc, meta.get("subtitle", ""), block_label="master draft")

    doc.add_heading("Статус сборки", level=1)
    doc.add_paragraph(
        "Draft: I-фаза заменена на птицеводство; блоки II–V и trade tables — as-is по плану "
        "(см. docs/teo-poultry/appendix/master-docx-assembly.md)."
    )
    doc.add_paragraph(f"Baseline DOCX: {meta.get('baseline_docx', '—')}")
    doc.add_paragraph(f"Статус: {meta.get('status', 'draft')}")

    for sec in assembly.get("sections", []):
        doc.add_page_break()
        heading = sec.get("title", sec.get("id", "?"))
        ref = sec.get("section_ref")
        if ref:
            heading = f"{heading} ({ref})"
        doc.add_heading(heading, level=1)

        action = sec.get("action", "replace")
        doc.add_paragraph(f"action: {action}")
        if sec.get("note"):
            p = doc.add_paragraph(sec["note"])
            for run in p.runs:
                run.italic = True

        refs = gather_section_sources(sec, registry)
        if refs:
            doc.add_paragraph("Источники:")
            for r in refs:
                add_ref_paragraph(doc, r.path or r.label)

        for r in refs:
            if not r.path:
                doc.add_paragraph(f"[{r.label}]")
                continue
            body = read_text(r.path)
            doc.add_heading(r.label, level=2)
            write_markdown_body(doc, body, r.path)

    doc.save(out)
    return out


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--skip-docx", action="store_true", help="Validate + audit only")
    parser.add_argument("--themes-only", action="store_true", help="Skip master assembly")
    parser.add_argument("--master-only", action="store_true", help="Only master draft")
    args = parser.parse_args()

    registry = load_registry()
    try:
        assert_registry_valid(registry)
    except Exception as exc:
        print(f"registry validation FAILED:\n{exc}", file=sys.stderr)
        return 1
    print("registry.yaml + 00-master-assembly.yaml OK")

    audit_path = write_docx_audit(registry)
    print(f"Wrote {audit_path}")

    if args.skip_docx:
        return 0

    paths: list[Path] = []
    if not args.master_only:
        paths.append(build_index_docx(registry))
        for tid in sorted(k for k in registry.get("themes", {}) if k.startswith("T")):
            paths.append(build_theme_docx(tid, registry["themes"][tid], registry))

    if not args.themes_only:
        assembly = load_master_assembly()
        paths.append(build_master_docx(registry, assembly))

    print(f"Generated {len(paths)} DOCX in {DOCX_DIR}")
    for p in paths:
        print(f"  {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
